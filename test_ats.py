import io
import docx
from parsers import (
    extract_text_from_pdf,
    extract_text_from_docx,
    parse_resume_sections,
    clean_text,
    scrape_job_description
)
from ats_engine import (
    analyze_job_description,
    calculate_ats_score,
    generate_ats_tailored_resume,
    generate_docx_resume
)

def test_document_extraction():
    print("[1] Testing DOCX extraction...")
    doc = docx.Document()
    doc.add_heading("Alex Rivera", level=1)
    doc.add_paragraph("alex.rivera@email.com | +1 555-0192 | Boston, MA")
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, React, TypeScript, SQL, Docker, AWS")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Software Engineering Intern at CloudTech (2024)")
    doc.add_paragraph("Engineered full-stack features using React and Python.")
    
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    extracted = extract_text_from_docx(docx_bytes)
    assert "Alex Rivera" in extracted, "Failed to extract name"
    assert "Python" in extracted, "Failed to extract skills"
    print("    -> DOCX extraction passed.")

    print("[2] Testing resume section parser...")
    parsed = parse_resume_sections(extracted)
    assert parsed["full_name"] == "Alex Rivera", f"Expected Alex Rivera, got {parsed['full_name']}"
    assert parsed["email"] == "alex.rivera@email.com", f"Expected email, got {parsed['email']}"
    assert "Python" in parsed["skills"] or "python" in [s.lower() for s in parsed["skills"]], "Expected Python in skills"
    print("    -> Resume section parsing passed.")

def test_ats_analysis_and_generation():
    print("[3] Testing JD analysis...")
    sample_jd = """Role: Senior Python Full Stack Engineer
Required Skills:
- 3+ years experience with Python, FastAPI, and Django.
- Strong proficiency in React, TypeScript, and Tailwind CSS.
- Hands-on experience with PostgreSQL, Redis, and Docker.
- Proven knowledge of AWS cloud services and CI/CD pipelines."""

    jd_analysis = analyze_job_description(sample_jd)
    assert "python" in jd_analysis["skills"], "Expected python in JD skills"
    assert "react" in jd_analysis["skills"], "Expected react in JD skills"
    assert "postgresql" in jd_analysis["skills"], "Expected postgresql in JD skills"
    print(f"    -> Extracted {len(jd_analysis['skills'])} skills from JD: {jd_analysis['skills']}")

    print("[4] Testing ATS resume generation and scoring...")
    student_data = {
        "full_name": "Jordan Lee",
        "email": "jordan.lee@cs.edu",
        "phone": "+1 555-987-6543",
        "location": "San Francisco, CA",
        "target_role": "Python Full Stack Engineer",
        "summary": "Motivated developer with experience in Python and web technologies.",
        "skills": ["Python", "FastAPI", "React", "PostgreSQL", "Git", "Docker"],
        "education": ["B.S. in Computer Science | UC Berkeley (2025)"],
        "experience": [
            "Software Intern | DevStudio (2024)\n- Built backend REST APIs using Python and PostgreSQL.\n- Collaborated with frontend engineers to integrate React components."
        ],
        "projects": [
            "TaskFlow Web Platform | React, FastAPI, PostgreSQL\n- Created task management app with real-time updates and database indexing."
        ]
    }

    tailored = generate_ats_tailored_resume(student_data, sample_jd)
    ats_score = tailored["ats_score"]
    
    print(f"    -> Overall ATS Score: {ats_score['overall_score']}% ({ats_score['rating']})")
    print(f"    -> Keyword match: {ats_score['breakdown']['keyword_match']}%")
    print(f"    -> Skills alignment: {ats_score['breakdown']['skills_alignment']}%")
    print(f"    -> Matched keywords: {ats_score['matched_keywords']}")
    print(f"    -> Missing keywords: {ats_score['missing_keywords']}")
    
    assert ats_score["overall_score"] >= 60, "Expected competitive ATS score"
    assert len(ats_score["matched_keywords"]) > 0, "Expected matched keywords"
    assert len(tailored["experience"][0]["bullets"]) > 0, "Expected tailored experience bullets"

    print("[5] Testing DOCX export generation...")
    docx_buf = generate_docx_resume(tailored)
    assert docx_buf.getbuffer().nbytes > 1000, "Expected non-empty docx binary"
    print("    -> DOCX generation passed.")

def test_url_scraping_resilience():
    print("[6] Testing URL scraper error handling...")
    res = scrape_job_description("https://invalid-non-existent-domain-test-12345.com/job")
    assert res["success"] is False, "Expected error on invalid domain"
    assert "Could not access URL" in res["error"]
    print("    -> Scraper error handling passed.")

if __name__ == "__main__":
    print("\n--- Running ATS Resume Generator Test Suite ---")
    test_document_extraction()
    test_ats_analysis_and_generation()
    test_url_scraping_resilience()
    print("\n[SUCCESS] All ATS tests passed successfully!\n")
