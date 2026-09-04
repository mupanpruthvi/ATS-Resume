import io
import re
from typing import Dict, Any, Optional
import requests
from bs4 import BeautifulSoup
import pypdf
import docx

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract clean text from a PDF file buffer using pypdf."""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    extracted = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            extracted.append(page_text.strip())
    raw_text = "\n\n".join(extracted)
    return clean_text(raw_text)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text and list items from a DOCX file buffer."""
    doc = docx.Document(io.BytesIO(file_bytes))
    lines = []
    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        if txt:
            lines.append(txt)
    for table in doc.tables:
        for row in table.rows:
            row_vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_vals:
                lines.append(" | ".join(row_vals))
    return clean_text("\n".join(lines))

def scrape_job_description(url: str) -> Dict[str, Any]:
    """
    Scrapes job posting content from a public URL.
    Returns cleaned title, raw text, and metadata.
    """
    url = url.strip()
    if not url.startswith("http://") and not url.startswith("https://"):
        url = "https://" + url

    try:
        response = requests.get(url, headers=HEADERS, timeout=12)
        response.raise_for_status()
    except Exception as exc:
        return {
            "success": False,
            "error": f"Could not access URL: {str(exc)}",
            "url": url,
            "text": "",
            "title": ""
        }

    soup = BeautifulSoup(response.text, "html.parser")

    # Remove script, style, header, nav, footer, noscript tags
    for tag in soup(["script", "style", "nav", "header", "footer", "aside", "svg", "noscript", "form", "iframe"]):
        tag.decompose()

    # Attempt to extract job title
    title = ""
    title_tags = [
        soup.find("h1"),
        soup.find("meta", property="og:title"),
        soup.find("title")
    ]
    for candidate in title_tags:
        if candidate:
            if candidate.name == "meta" and candidate.get("content"):
                title = candidate["content"].strip()
                break
            elif candidate.text and candidate.text.strip():
                title = candidate.text.strip()
                break

    # Look for common job posting containers
    content_container = None
    common_selectors = [
        '[class*="job-description"]',
        '[class*="jobDescription"]',
        '[class*="job_description"]',
        '[id*="job-description"]',
        '[id*="jobDescription"]',
        '[class*="description"]',
        '[class*="posting-requirements"]',
        '[class*="job-details"]',
        '[class*="content"]',
        "article",
        "main",
    ]

    for sel in common_selectors:
        found = soup.select_one(sel)
        if found and len(found.get_text(strip=True)) > 200:
            content_container = found
            break

    if content_container:
        body_text = content_container.get_text(separator="\n")
    else:
        body_tag = soup.find("body")
        body_text = body_tag.get_text(separator="\n") if body_tag else soup.get_text(separator="\n")

    cleaned_jd = clean_text(body_text)

    # Check if we retrieved meaningful content
    if len(cleaned_jd) < 100:
        return {
            "success": True,
            "warning": "Extracted text is brief. The website might require login (e.g. private LinkedIn post). You can paste the text directly.",
            "url": url,
            "title": title or "Job Description",
            "text": cleaned_jd
        }

    return {
        "success": True,
        "url": url,
        "title": title or "Job Description",
        "text": cleaned_jd
    }

def clean_text(text: str) -> str:
    """Normalize whitespace and remove non-printable characters."""
    if not text:
        return ""
    # Standardize unicode quotes and hyphens
    text = text.replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2013", "-").replace("\u2014", "-").replace("\u2022", "* ").replace("\u00a0", " ")
    # Replace carriage returns
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse multiple blank lines
    lines = [line.strip() for line in text.split("\n")]
    result_lines = []
    prev_empty = False
    for line in lines:
        if line:
            result_lines.append(line)
            prev_empty = False
        elif not prev_empty:
            result_lines.append("")
            prev_empty = True
    return "\n".join(result_lines).strip()

def parse_resume_sections(text: str) -> Dict[str, Any]:
    """
    Intelligently parses unstructured resume text into candidate profile fields.
    Extracts name, email, phone, links, summary, skills, education, projects, experience.
    """
    parsed: Dict[str, Any] = {
        "full_name": "",
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
        "target_role": "",
        "summary": "",
        "skills": [],
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "raw_text": text
    }

    if not text:
        return parsed

    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # Extract email
    email_match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    if email_match:
        parsed["email"] = email_match.group(0)

    # Extract phone
    phone_match = re.search(r"(\+?\d{1,3}[-.\s]?)?(\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}", text)
    if phone_match and len(phone_match.group(0).strip()) >= 10:
        parsed["phone"] = phone_match.group(0).strip()

    # Extract LinkedIn & GitHub
    linkedin_match = re.search(r"(?:https?:\/\/)?(?:www\.)?linkedin\.com\/in\/[a-zA-Z0-9_-]+", text, re.I)
    if linkedin_match:
        parsed["linkedin"] = linkedin_match.group(0)

    github_match = re.search(r"(?:https?:\/\/)?(?:www\.)?github\.com\/[a-zA-Z0-9_-]+", text, re.I)
    if github_match:
        parsed["github"] = github_match.group(0)

    # Candidate Name (typically in top 5 lines, before email/phone)
    for line in lines[:5]:
        if parsed["email"] and parsed["email"] in line:
            continue
        if parsed["phone"] and parsed["phone"] in line:
            continue
        if "curriculum vitae" in line.lower() or "resume" in line.lower():
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(re.match(r"^[A-Za-z.\-']+$", w) for w in words):
            parsed["full_name"] = line
            break

    if not parsed["full_name"] and lines:
        parsed["full_name"] = lines[0][:40]

    # Section boundaries
    section_patterns = {
        "summary": r"(?:summary|objective|professional summary|about me|profile)",
        "skills": r"(?:technical skills|skills & tools|skills|core competencies|technologies|expertise)",
        "experience": r"(?:work experience|professional experience|experience|employment history|internships|work history)",
        "projects": r"(?:projects|key projects|academic projects|personal projects)",
        "education": r"(?:education|academic background|qualifications|academic history)",
        "certifications": r"(?:certifications|certificates|licenses|courses|awards|achievements)"
    }

    current_section = None
    section_buffers: Dict[str, list] = {k: [] for k in section_patterns}

    for line in lines:
        line_lower = line.lower().strip(":").strip()
        matched_section = None
        for sec, pat in section_patterns.items():
            if re.fullmatch(pat, line_lower, re.I) or re.match(r"^#+\s*" + pat + "$", line_lower, re.I):
                matched_section = sec
                break

        if matched_section:
            current_section = matched_section
        elif current_section:
            section_buffers[current_section].append(line)
        else:
            if not parsed["target_role"] and line != parsed["full_name"] and len(line) < 60:
                if not any(c in line for c in ["@", "http", "+", "linkedin", "github"]):
                    parsed["target_role"] = line

    if section_buffers["summary"]:
        parsed["summary"] = " ".join(section_buffers["summary"])

    if section_buffers["skills"]:
        skill_text = " ".join(section_buffers["skills"])
        raw_skills = re.split(r"[,;|•\*\n]+", skill_text)
        cleaned_skills = []
        for s in raw_skills:
            item = s.strip().strip("-").strip()
            if ":" in item:
                parts = item.split(":", 1)
                item = parts[1].strip()
            if item and len(item) < 40 and not any(w in item.lower() for w in ["including", "proficient in", "knowledge of"]):
                cleaned_skills.append(item)
        parsed["skills"] = list(dict.fromkeys(cleaned_skills))

    if section_buffers["education"]:
        edu_entries = []
        cur_entry = []
        for l in section_buffers["education"]:
            if any(term in l.lower() for term in ["bachelor", "master", "b.tech", "m.tech", "b.s.", "m.s.", "diploma", "university", "college", "institute", "high school"]):
                if cur_entry:
                    edu_entries.append("\n".join(cur_entry))
                    cur_entry = []
            cur_entry.append(l)
        if cur_entry:
            edu_entries.append("\n".join(cur_entry))
        parsed["education"] = edu_entries if edu_entries else ["\n".join(section_buffers["education"])]

    if section_buffers["experience"]:
        exp_entries = []
        cur_entry = []
        for l in section_buffers["experience"]:
            if re.search(r"\b(20\d\d|19\d\d|present|current|intern|engineer|developer|analyst)\b", l, re.I) and not l.startswith(("-", "*", "•")):
                if len(cur_entry) > 1:
                    exp_entries.append("\n".join(cur_entry))
                    cur_entry = []
            cur_entry.append(l)
        if cur_entry:
            exp_entries.append("\n".join(cur_entry))
        parsed["experience"] = exp_entries if exp_entries else ["\n".join(section_buffers["experience"])]

    if section_buffers["projects"]:
        proj_entries = []
        cur_proj = []
        for l in section_buffers["projects"]:
            if not l.startswith(("-", "*", "•")) and (len(l) < 60 or "|" in l):
                if cur_proj:
                    proj_entries.append("\n".join(cur_proj))
                    cur_proj = []
            cur_proj.append(l)
        if cur_proj:
            proj_entries.append("\n".join(cur_proj))
        parsed["projects"] = proj_entries if proj_entries else ["\n".join(section_buffers["projects"])]

    if section_buffers["certifications"]:
        parsed["certifications"] = [c.strip().lstrip("•*- ") for c in section_buffers["certifications"] if c.strip()]

    return parsed
